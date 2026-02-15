"""
Document Structure Extraction Service - FEATURE #134

Extracts structured content from documents (PDF, Word, Markdown) while preserving:
- Headings (with levels)
- Paragraphs
- Lists (bullet and numbered)
- Tables
- Code blocks (for Markdown)

This structured representation is then passed to the semantic chunker for better
chunking that respects document structure.
"""

import re
import logging
from typing import List, Dict, Optional
from pathlib import Path
from enum import Enum

logger = logging.getLogger(__name__)


class ElementType(str, Enum):
    """Types of document elements."""
    HEADING = "heading"
    PARAGRAPH = "paragraph"
    LIST = "list"
    TABLE = "table"
    CODE = "code"


class DocumentElement:
    """Represents a single structural element in a document."""

    def __init__(
        self,
        element_type: ElementType,
        content: str,
        level: Optional[int] = None,
        metadata: Optional[Dict] = None
    ):
        self.type = element_type
        self.content = content
        self.level = level  # For headings: 1-6, for lists: nesting level
        self.metadata = metadata or {}

    def to_dict(self) -> Dict:
        """Convert element to dictionary representation."""
        return {
            "type": self.type.value,
            "content": self.content,
            "level": self.level,
            "metadata": self.metadata
        }


class DocumentStructureExtractor:
    """
    Extracts structured content from various document formats.

    Supported formats:
    - PDF: Uses pdfplumber (if available) or pypdf as fallback
    - Word (.docx): Uses python-docx with heading style detection
    - Markdown: Parses heading levels, code blocks, and lists
    - Plain text: Detects headings heuristically
    """

    def __init__(self):
        # Check which PDF libraries are available
        self.pdfplumber_available = self._check_pdfplumber()
        self.pymupdf_available = self._check_pymupdf()

        if self.pdfplumber_available:
            logger.info("Using pdfplumber for PDF structure extraction")
        elif self.pymupdf_available:
            logger.info("Using pymupdf (fitz) for PDF structure extraction")
        else:
            logger.info("Using pypdf as fallback for PDF extraction (structure limited)")

    def _check_pdfplumber(self) -> bool:
        """Check if pdfplumber is available."""
        try:
            import pdfplumber
            return True
        except ImportError:
            return False

    def _check_pymupdf(self) -> bool:
        """Check if pymupdf (fitz) is available."""
        try:
            import fitz  # pymupdf
            return True
        except ImportError:
            return False

    def extract_structure(self, file_path: Path, mime_type: str) -> List[DocumentElement]:
        """
        Extract structured elements from a document.

        Args:
            file_path: Path to the document file
            mime_type: MIME type of the document

        Returns:
            List of DocumentElement objects representing the document structure
        """
        try:
            if mime_type == "application/pdf":
                return self._extract_pdf_structure(file_path)
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self._extract_word_structure(file_path)
            elif mime_type == "text/markdown":
                return self._extract_markdown_structure(file_path)
            elif mime_type == "text/plain":
                return self._extract_text_structure(file_path)
            else:
                logger.warning(f"Unsupported mime type for structure extraction: {mime_type}")
                return []
        except Exception as e:
            logger.error(f"Error extracting document structure: {e}")
            return []

    def _extract_pdf_structure(self, file_path: Path) -> List[DocumentElement]:
        """
        Extract structure from PDF using pdfplumber (preferred) or pymupdf (fallback).

        PDFplumber can detect:
        - Text with font information (size, bold, etc.)
        - Tables
        - Page layout

        PyMuPDF can detect:
        - Text blocks with font information
        - Tables (limited)
        - Headings based on font size
        """
        if self.pdfplumber_available:
            return self._extract_pdf_with_pdfplumber(file_path)
        elif self.pymupdf_available:
            return self._extract_pdf_with_pymupdf(file_path)
        else:
            return self._extract_pdf_with_pypdf(file_path)

    def _extract_pdf_with_pdfplumber(self, file_path: Path) -> List[DocumentElement]:
        """Extract PDF structure using pdfplumber."""
        try:
            import pdfplumber

            elements = []

            with pdfplumber.open(str(file_path)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text with layout information
                    words = page.extract_words(
                        x_tolerance=3,
                        y_tolerance=3,
                        keep_blank_chars=False
                    )

                    if not words:
                        continue

                    # Group words into lines based on y-position
                    lines = self._group_words_into_lines(words)

                    # Detect headings based on font size
                    # Calculate average font size for the page
                    avg_height = sum(w.get('height', 12) for w in words) / len(words) if words else 12

                    for line in lines:
                        text = line['text'].strip()
                        if not text:
                            continue

                        height = line.get('height', avg_height)
                        is_bold = any(w.get('chars', [{}])[0].get('fontname', '').lower().find('bold') >= 0
                                     for w in line.get('words', []) if w.get('chars'))

                        # Detect heading based on size and bold
                        if height > avg_height * 1.3 or (height > avg_height * 1.1 and is_bold):
                            # Determine heading level based on size
                            if height > avg_height * 1.8:
                                level = 1
                            elif height > avg_height * 1.5:
                                level = 2
                            else:
                                level = 3

                            elements.append(DocumentElement(
                                element_type=ElementType.HEADING,
                                content=text,
                                level=level,
                                metadata={'page': page_num, 'font_size': height}
                            ))
                        else:
                            # Regular paragraph
                            elements.append(DocumentElement(
                                element_type=ElementType.PARAGRAPH,
                                content=text,
                                metadata={'page': page_num}
                            ))

                    # Extract tables
                    # FEATURE #215: Enhanced table extraction with smart merging
                    # PDFs often have headers and data in separate "tables" due to layout
                    tables = page.extract_tables()
                    merged_tables = self._merge_related_tables(tables)

                    for table_idx, table in enumerate(merged_tables):
                        if table and len(table) > 0:
                            # Convert table to text representation with page number for context
                            table_text = self._table_to_text(table, page_num=page_num)
                            if table_text.strip():  # Only add non-empty tables
                                # Detect if it's a price table for metadata
                                is_price = self._detect_price_table([[str(c) if c else "" for c in row] for row in table if row])
                                elements.append(DocumentElement(
                                    element_type=ElementType.TABLE,
                                    content=table_text,
                                    metadata={
                                        'page': page_num,
                                        'table_index': table_idx,
                                        'is_price_table': is_price,
                                        'row_count': len(table)
                                    }
                                ))

            logger.info(f"Extracted {len(elements)} structured elements from PDF using pdfplumber")
            return elements

        except Exception as e:
            logger.error(f"Error extracting PDF with pdfplumber: {e}")
            # Fallback to pymupdf or pypdf
            if self.pymupdf_available:
                return self._extract_pdf_with_pymupdf(file_path)
            else:
                return self._extract_pdf_with_pypdf(file_path)

    def _merge_related_tables(self, tables: List[List]) -> List[List]:
        """
        FEATURE #215: Merge related tables that share the same column structure.

        PDFs often extract what appears as one table as multiple separate tables
        due to visual separations (lines, spacing). This method merges tables
        that have compatible column structures.

        Strategy:
        1. Find header tables (single row with column names like Model, Price, etc.)
        2. Merge subsequent data tables that match the column count
        3. Keep non-price tables separate
        """
        if not tables:
            return []

        merged = []
        current_header = None
        current_merged = []

        for table in tables:
            if not table or not any(any(cell for cell in row) for row in table):
                continue

            # Clean the table
            cleaned = [[str(c).strip() if c else "" for c in row] for row in table]

            # Check if this looks like a header row
            if len(cleaned) == 1:
                first_row = cleaned[0]
                # Check for price table header keywords
                row_text = ' '.join(first_row).lower()
                is_header = any(kw in row_text for kw in ['model', 'part', 'description', 'price', 'cost', 'availability'])

                if is_header:
                    # Finish any previous merged table
                    if current_merged:
                        merged.append(current_merged)

                    # Start new merged table with this header
                    current_header = first_row
                    current_merged = [first_row]
                    continue

            # Check if this table matches current header's column count
            if current_header and len(current_header) > 0:
                # Tables with same or compatible column count can be merged
                for row in cleaned:
                    if len(row) == len(current_header) or (len(row) > 0 and len(row) <= len(current_header)):
                        # Pad row if needed
                        while len(row) < len(current_header):
                            row.append("")
                        current_merged.append(row)
            else:
                # No current header - treat as standalone table
                if current_merged:
                    merged.append(current_merged)
                    current_merged = []
                    current_header = None

                # Check if first row looks like a header for this table
                if cleaned and len(cleaned) > 1:
                    first_row_text = ' '.join(cleaned[0]).lower()
                    if any(kw in first_row_text for kw in ['model', 'part', 'description', 'price']):
                        current_header = cleaned[0]
                        current_merged = cleaned
                    else:
                        merged.append(cleaned)
                else:
                    merged.append(cleaned)

        # Don't forget the last merged table
        if current_merged:
            merged.append(current_merged)

        return merged

    def _group_words_into_lines(self, words: List[Dict]) -> List[Dict]:
        """Group words into lines based on y-position."""
        if not words:
            return []

        # Sort words by y-position (top to bottom)
        sorted_words = sorted(words, key=lambda w: (w.get('top', 0), w.get('x0', 0)))

        lines = []
        current_line = {
            'text': '',
            'words': [],
            'top': sorted_words[0].get('top', 0),
            'height': sorted_words[0].get('height', 12)
        }

        for word in sorted_words:
            y_pos = word.get('top', 0)

            # If word is on the same line (within 2 pixels)
            if abs(y_pos - current_line['top']) < 2:
                current_line['text'] += ' ' + word.get('text', '')
                current_line['words'].append(word)
                current_line['height'] = max(current_line['height'], word.get('height', 12))
            else:
                # New line
                if current_line['text'].strip():
                    lines.append(current_line)
                current_line = {
                    'text': word.get('text', ''),
                    'words': [word],
                    'top': y_pos,
                    'height': word.get('height', 12)
                }

        # Add last line
        if current_line['text'].strip():
            lines.append(current_line)

        return lines

    def _table_to_text(self, table: List[List], page_num: int = 0) -> str:
        """
        Convert table data to text representation.

        FEATURE #215: Enhanced table-to-text conversion for better RAG retrieval.
        Creates a semantic format that preserves the relationship between columns,
        especially for price lists and product catalogs.

        Args:
            table: 2D list of table cells
            page_num: Page number (1-indexed) for context

        Returns:
            Formatted text representation of the table
        """
        if not table:
            return ""

        # Clean table data - remove empty rows
        cleaned_table = []
        for row in table:
            if row and any(cell and str(cell).strip() for cell in row):
                cleaned_row = [str(cell).strip() if cell else "" for cell in row]
                cleaned_table.append(cleaned_row)

        if not cleaned_table:
            return ""

        # Try to detect if this is a price/product table
        is_price_table = self._detect_price_table(cleaned_table)

        if is_price_table:
            return self._format_price_table(cleaned_table, page_num)
        else:
            return self._format_generic_table(cleaned_table, page_num)

    def _detect_price_table(self, table: List[List[str]]) -> bool:
        """
        Detect if a table contains price data.

        Looks for common price table indicators:
        - Price/Cost/EUR/USD/List Price column headers
        - Currency symbols or numeric values with decimal/comma
        - Product/Part/Model number columns
        """
        if not table:
            return False

        # Check first row (likely header) for price indicators
        header = ' '.join(str(cell).lower() for cell in table[0] if cell)
        price_indicators = ['price', 'cost', 'eur', 'usd', 'gbp', '€', '$', '£', 'amount', 'total', 'list']
        product_indicators = ['model', 'part', 'product', 'item', 'code', 'sku', 'number']

        has_price_col = any(ind in header for ind in price_indicators)
        has_product_col = any(ind in header for ind in product_indicators)

        if has_price_col and has_product_col:
            return True

        # Check data rows for numeric patterns that look like prices
        price_pattern = re.compile(r'^[\d,.\s]+$')  # Numbers with commas/decimals
        for row in table[1:4]:  # Check first few data rows
            for cell in row:
                if cell and price_pattern.match(cell.replace(' ', '')):
                    # Looks like a price value
                    try:
                        # Try to parse as number
                        val = float(cell.replace(',', '').replace(' ', ''))
                        if 10 < val < 10000000:  # Reasonable price range
                            return True
                    except ValueError:
                        pass

        return False

    def _format_price_table(self, table: List[List[str]], page_num: int) -> str:
        """
        Format a price table in a semantic way for optimal RAG retrieval.

        FEATURE #215: Creates natural language entries for each product/price row,
        making it easy for LLMs to find and report prices accurately.

        Example output:
        "Product: VFR-X1M06SA-AAA-AA9 | Description: VMFT X-Band 10kW Masthead (6ft) | Price: 13,000 EUR | Availability: STD"
        """
        if len(table) < 2:
            return self._format_generic_table(table, page_num)

        lines = []
        header = table[0]

        # Identify column indices for key fields
        col_indices = self._identify_price_table_columns(header)

        # Add header info
        lines.append(f"[Price Table - Page {page_num}]")
        lines.append(f"Columns: {' | '.join(h for h in header if h)}")
        lines.append("")

        # Format each data row
        for row_idx, row in enumerate(table[1:], 1):
            if not any(cell for cell in row):
                continue

            parts = []

            # Build semantic entry
            if col_indices.get('model') is not None and row[col_indices['model']]:
                parts.append(f"Model/Part: {row[col_indices['model']]}")

            if col_indices.get('description') is not None and row[col_indices['description']]:
                parts.append(f"Description: {row[col_indices['description']]}")

            if col_indices.get('price') is not None and row[col_indices['price']]:
                price_val = row[col_indices['price']]
                # Add currency if not already present
                if not any(c in price_val for c in ['€', '$', '£', 'EUR', 'USD', 'GBP']):
                    price_val = f"{price_val} EUR"  # Default to EUR for this price book
                parts.append(f"Price: {price_val}")

            if col_indices.get('availability') is not None and row[col_indices['availability']]:
                parts.append(f"Availability: {row[col_indices['availability']]}")

            # Add any remaining columns
            for idx, cell in enumerate(row):
                if cell and idx not in col_indices.values():
                    col_name = header[idx] if idx < len(header) and header[idx] else f"Col{idx}"
                    parts.append(f"{col_name}: {cell}")

            if parts:
                lines.append(" | ".join(parts))

        return "\n".join(lines)

    def _identify_price_table_columns(self, header: List[str]) -> Dict[str, int]:
        """Identify column indices for common price table fields."""
        indices = {}

        for idx, col in enumerate(header):
            if not col:
                continue
            col_lower = col.lower()

            # Model/Part Number
            if any(kw in col_lower for kw in ['model', 'part', 'product', 'item', 'code', 'sku']):
                if 'model' not in indices:  # Take first match
                    indices['model'] = idx

            # Description
            elif any(kw in col_lower for kw in ['description', 'desc', 'name', 'details']):
                if 'description' not in indices:
                    indices['description'] = idx

            # Price
            elif any(kw in col_lower for kw in ['price', 'cost', 'eur', 'usd', 'amount', 'total', '€', '$']):
                if 'price' not in indices:
                    indices['price'] = idx

            # Availability
            elif any(kw in col_lower for kw in ['availability', 'avail', 'stock', 'status', 'lead']):
                if 'availability' not in indices:
                    indices['availability'] = idx

        return indices

    def _format_generic_table(self, table: List[List[str]], page_num: int) -> str:
        """
        Format a generic table with header row and data rows.

        Uses markdown-style table format for readability.
        """
        if not table:
            return ""

        lines = []
        lines.append(f"[Table - Page {page_num}]")

        # Add header row
        header = table[0]
        lines.append(" | ".join(h if h else "" for h in header))

        # Add separator
        lines.append(" | ".join("-" * max(len(h), 3) if h else "---" for h in header))

        # Add data rows
        for row in table[1:]:
            if any(cell for cell in row):
                row_text = " | ".join(str(cell) if cell else "" for cell in row)
                lines.append(row_text)

        return "\n".join(lines)

    def _extract_pdf_with_pymupdf(self, file_path: Path) -> List[DocumentElement]:
        """Extract PDF structure using pymupdf (fitz)."""
        try:
            import fitz  # pymupdf

            elements = []

            doc = fitz.open(str(file_path))

            for page_num in range(len(doc)):
                page = doc[page_num]

                # Get text blocks with formatting info
                blocks = page.get_text("dict")["blocks"]

                avg_size = 12.0  # Default
                sizes = []

                # Collect all font sizes to determine average
                for block in blocks:
                    if block.get("type") == 0:  # Text block
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                sizes.append(span.get("size", 12))

                if sizes:
                    avg_size = sum(sizes) / len(sizes)

                # Process blocks
                for block in blocks:
                    if block.get("type") == 0:  # Text block
                        for line in block.get("lines", []):
                            line_text = ""
                            max_size = avg_size
                            is_bold = False

                            for span in line.get("spans", []):
                                line_text += span.get("text", "")
                                span_size = span.get("size", avg_size)
                                max_size = max(max_size, span_size)

                                # Check if bold
                                font = span.get("font", "").lower()
                                if "bold" in font:
                                    is_bold = True

                            line_text = line_text.strip()
                            if not line_text:
                                continue

                            # Detect heading
                            if max_size > avg_size * 1.3 or (max_size > avg_size * 1.1 and is_bold):
                                if max_size > avg_size * 1.8:
                                    level = 1
                                elif max_size > avg_size * 1.5:
                                    level = 2
                                else:
                                    level = 3

                                elements.append(DocumentElement(
                                    element_type=ElementType.HEADING,
                                    content=line_text,
                                    level=level,
                                    metadata={'page': page_num + 1, 'font_size': max_size}
                                ))
                            else:
                                elements.append(DocumentElement(
                                    element_type=ElementType.PARAGRAPH,
                                    content=line_text,
                                    metadata={'page': page_num + 1}
                                ))

            doc.close()

            logger.info(f"Extracted {len(elements)} structured elements from PDF using pymupdf")
            return elements

        except Exception as e:
            logger.error(f"Error extracting PDF with pymupdf: {e}")
            return self._extract_pdf_with_pypdf(file_path)

    def _extract_pdf_with_pypdf(self, file_path: Path) -> List[DocumentElement]:
        """
        Extract PDF structure using pypdf (basic fallback).

        pypdf doesn't preserve font information, so we use heuristics:
        - ALL CAPS lines -> headings
        - Short lines followed by blank -> headings
        - Everything else -> paragraphs
        """
        try:
            from pypdf import PdfReader

            elements = []
            reader = PdfReader(str(file_path))

            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if not text:
                    continue

                lines = text.split('\n')

                for line in lines:
                    line = line.strip()
                    if not line:
                        continue

                    # Heuristic: ALL CAPS and short -> heading
                    if len(line) < 80 and line.isupper() and len(line.split()) >= 2:
                        elements.append(DocumentElement(
                            element_type=ElementType.HEADING,
                            content=line,
                            level=2,  # Default level
                            metadata={'page': page_num, 'detection': 'all_caps'}
                        ))
                    # Heuristic: Numbered section (1. or 1.1)
                    elif re.match(r'^\d+\.(?:\d+\.)*\s+', line):
                        elements.append(DocumentElement(
                            element_type=ElementType.HEADING,
                            content=line,
                            level=2,
                            metadata={'page': page_num, 'detection': 'numbered'}
                        ))
                    else:
                        elements.append(DocumentElement(
                            element_type=ElementType.PARAGRAPH,
                            content=line,
                            metadata={'page': page_num}
                        ))

            logger.info(f"Extracted {len(elements)} elements from PDF using pypdf (limited structure)")
            return elements

        except Exception as e:
            logger.error(f"Error extracting PDF with pypdf: {e}")
            return []

    def _extract_word_structure(self, file_path: Path) -> List[DocumentElement]:
        """
        Extract structure from Word document using python-docx.

        Detects:
        - Heading styles (Heading 1, 2, 3, etc.)
        - Lists (bullet and numbered)
        - Tables
        - Regular paragraphs
        """
        try:
            from docx import Document as DocxDocument
            from docx.enum.text import WD_STYLE_TYPE

            elements = []
            doc = DocxDocument(str(file_path))

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else "Normal"

                # Detect headings by style
                if style_name.startswith("Heading"):
                    # Extract level from style name (e.g., "Heading 1" -> 1)
                    level_match = re.search(r'Heading\s+(\d+)', style_name)
                    level = int(level_match.group(1)) if level_match else 1

                    elements.append(DocumentElement(
                        element_type=ElementType.HEADING,
                        content=text,
                        level=level,
                        metadata={'style': style_name}
                    ))
                elif style_name.startswith("List"):
                    # List item
                    elements.append(DocumentElement(
                        element_type=ElementType.LIST,
                        content=text,
                        metadata={'style': style_name}
                    ))
                else:
                    # Regular paragraph
                    elements.append(DocumentElement(
                        element_type=ElementType.PARAGRAPH,
                        content=text,
                        metadata={'style': style_name}
                    ))

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    table_text.append(row_text)

                if table_text:
                    elements.append(DocumentElement(
                        element_type=ElementType.TABLE,
                        content="\n".join(table_text),
                        metadata={'table_index': table_idx, 'rows': len(table.rows)}
                    ))

            logger.info(f"Extracted {len(elements)} structured elements from Word document")
            return elements

        except Exception as e:
            logger.error(f"Error extracting Word structure: {e}")
            return []

    def _extract_markdown_structure(self, file_path: Path) -> List[DocumentElement]:
        """
        Extract structure from Markdown file.

        Detects:
        - Headings (#, ##, ###, etc.)
        - Code blocks (```...```)
        - Lists (- or * or numbered)
        - Tables (| ... |)
        - Paragraphs
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            elements = []
            lines = content.split('\n')
            i = 0

            while i < len(lines):
                line = lines[i]
                stripped = line.strip()

                if not stripped:
                    i += 1
                    continue

                # Heading
                if stripped.startswith('#'):
                    level = len(re.match(r'^(#+)', stripped).group(1))
                    heading_text = re.sub(r'^#+\s*', '', stripped)
                    elements.append(DocumentElement(
                        element_type=ElementType.HEADING,
                        content=heading_text,
                        level=level,
                        metadata={'markdown': True}
                    ))
                    i += 1

                # Code block
                elif stripped.startswith('```'):
                    code_lines = [stripped[3:]]  # Language identifier
                    i += 1

                    while i < len(lines) and not lines[i].strip().startswith('```'):
                        code_lines.append(lines[i])
                        i += 1

                    if i < len(lines):  # Found closing ```
                        i += 1

                    code_content = '\n'.join(code_lines)
                    elements.append(DocumentElement(
                        element_type=ElementType.CODE,
                        content=code_content,
                        metadata={'language': code_lines[0].strip() if code_lines else 'text'}
                    ))

                # List (bullet or numbered)
                elif re.match(r'^[-*]\s+', stripped) or re.match(r'^\d+\.\s+', stripped):
                    list_lines = [stripped]
                    i += 1

                    # Collect consecutive list items
                    while i < len(lines):
                        next_line = lines[i].strip()
                        if re.match(r'^[-*]\s+', next_line) or re.match(r'^\d+\.\s+', next_line):
                            list_lines.append(next_line)
                            i += 1
                        elif not next_line:  # Empty line ends list
                            i += 1
                            break
                        else:
                            break

                    list_content = '\n'.join(list_lines)
                    elements.append(DocumentElement(
                        element_type=ElementType.LIST,
                        content=list_content,
                        metadata={'items': len(list_lines)}
                    ))

                # Table
                elif '|' in stripped and stripped.count('|') >= 2:
                    table_lines = [stripped]
                    i += 1

                    # Collect consecutive table rows
                    while i < len(lines):
                        next_line = lines[i].strip()
                        if '|' in next_line:
                            table_lines.append(next_line)
                            i += 1
                        elif not next_line:
                            i += 1
                            break
                        else:
                            break

                    table_content = '\n'.join(table_lines)
                    elements.append(DocumentElement(
                        element_type=ElementType.TABLE,
                        content=table_content,
                        metadata={'rows': len(table_lines)}
                    ))

                # Paragraph
                else:
                    para_lines = [stripped]
                    i += 1

                    # Collect consecutive non-empty lines
                    while i < len(lines):
                        next_line = lines[i].strip()
                        if not next_line:  # Empty line ends paragraph
                            i += 1
                            break
                        # Check if next line is a special element
                        if (next_line.startswith('#') or
                            next_line.startswith('```') or
                            re.match(r'^[-*]\s+', next_line) or
                            re.match(r'^\d+\.\s+', next_line) or
                            '|' in next_line):
                            break
                        para_lines.append(next_line)
                        i += 1

                    para_content = ' '.join(para_lines)
                    elements.append(DocumentElement(
                        element_type=ElementType.PARAGRAPH,
                        content=para_content
                    ))

            logger.info(f"Extracted {len(elements)} structured elements from Markdown")
            return elements

        except Exception as e:
            logger.error(f"Error extracting Markdown structure: {e}")
            return []

    def _extract_text_structure(self, file_path: Path) -> List[DocumentElement]:
        """
        Extract structure from plain text file using heuristics.

        Detects:
        - ALL CAPS lines as headings
        - Numbered sections as headings
        - Empty lines as paragraph breaks
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            elements = []
            paragraphs = content.split('\n\n')  # Split on double newlines

            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue

                lines = para.split('\n')
                first_line = lines[0].strip()

                # Check if it's a heading (ALL CAPS or numbered)
                if len(first_line) < 80 and first_line.isupper() and len(first_line.split()) >= 2:
                    elements.append(DocumentElement(
                        element_type=ElementType.HEADING,
                        content=para,
                        level=2,
                        metadata={'detection': 'all_caps'}
                    ))
                elif re.match(r'^\d+\.(?:\d+\.)*\s+', first_line):
                    elements.append(DocumentElement(
                        element_type=ElementType.HEADING,
                        content=para,
                        level=2,
                        metadata={'detection': 'numbered'}
                    ))
                else:
                    elements.append(DocumentElement(
                        element_type=ElementType.PARAGRAPH,
                        content=para
                    ))

            logger.info(f"Extracted {len(elements)} elements from plain text")
            return elements

        except Exception as e:
            logger.error(f"Error extracting text structure: {e}")
            return []

    # ==========================================================================
    # FEATURE #216: Hybrid PDF processing (text + tables)
    # ==========================================================================

    def extract_hybrid_pdf(self, file_path: Path) -> Dict:
        """
        FEATURE #216: Extract both structured table data AND text from a PDF.

        This method processes PDFs with mixed content by:
        1. Detecting table regions and extracting them as structured rows for SQL queries
        2. Extracting text regions for semantic chunking and vector embeddings

        Args:
            file_path: Path to the PDF file

        Returns:
            Dict with:
                - 'table_data': List of dicts with 'schema' and 'rows' for each table
                - 'text_elements': List of DocumentElement for text content
                - 'has_tables': Boolean indicating if tables were found
                - 'table_count': Number of tables extracted
                - 'row_count': Total number of rows across all tables
        """
        result = {
            'table_data': [],
            'text_elements': [],
            'has_tables': False,
            'table_count': 0,
            'row_count': 0
        }

        if not self.pdfplumber_available:
            logger.warning("pdfplumber not available - falling back to text-only extraction")
            result['text_elements'] = self._extract_pdf_structure(file_path)
            return result

        try:
            import pdfplumber

            text_elements = []
            all_table_rows = []
            unified_schema = None

            with pdfplumber.open(str(file_path)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # ===== EXTRACT TEXT REGIONS =====
                    words = page.extract_words(
                        x_tolerance=3,
                        y_tolerance=3,
                        keep_blank_chars=False
                    )

                    if words:
                        # Group words into lines
                        lines = self._group_words_into_lines(words)
                        avg_height = sum(w.get('height', 12) for w in words) / len(words) if words else 12

                        for line in lines:
                            text = line['text'].strip()
                            if not text:
                                continue

                            height = line.get('height', avg_height)
                            is_bold = any(
                                w.get('chars', [{}])[0].get('fontname', '').lower().find('bold') >= 0
                                for w in line.get('words', []) if w.get('chars')
                            )

                            # Detect heading
                            if height > avg_height * 1.3 or (height > avg_height * 1.1 and is_bold):
                                level = 1 if height > avg_height * 1.8 else (2 if height > avg_height * 1.5 else 3)
                                text_elements.append(DocumentElement(
                                    element_type=ElementType.HEADING,
                                    content=text,
                                    level=level,
                                    metadata={'page': page_num, 'font_size': height}
                                ))
                            else:
                                text_elements.append(DocumentElement(
                                    element_type=ElementType.PARAGRAPH,
                                    content=text,
                                    metadata={'page': page_num}
                                ))

                    # ===== EXTRACT TABLE REGIONS FOR STRUCTURED STORAGE =====
                    tables = page.extract_tables()
                    if tables:
                        merged_tables = self._merge_related_tables(tables)

                        for table_idx, table in enumerate(merged_tables):
                            if not table or len(table) < 2:  # Need at least header + 1 data row
                                continue

                            # Clean the table
                            cleaned = [[str(c).strip() if c else "" for c in row] for row in table]

                            # First row is assumed to be header
                            header = cleaned[0]
                            if not any(h for h in header):
                                continue

                            # Check if this looks like a data table (has price/product columns)
                            is_price_table = self._detect_price_table(cleaned)

                            # Extract rows as dicts for structured storage
                            for row in cleaned[1:]:
                                if not any(cell for cell in row):
                                    continue

                                # Create dict with header keys
                                row_dict = {}
                                for col_idx, cell in enumerate(row):
                                    if col_idx < len(header) and header[col_idx]:
                                        row_dict[header[col_idx]] = cell
                                    elif cell:
                                        row_dict[f'Column_{col_idx + 1}'] = cell

                                if row_dict:
                                    all_table_rows.append(row_dict)

                            # Track unified schema
                            if unified_schema is None:
                                unified_schema = [h for h in header if h]
                            else:
                                # Merge schemas
                                for h in header:
                                    if h and h not in unified_schema:
                                        unified_schema.append(h)

                            # Also add table as text element for embedding context
                            table_text = self._table_to_text(cleaned, page_num=page_num)
                            if table_text.strip():
                                text_elements.append(DocumentElement(
                                    element_type=ElementType.TABLE,
                                    content=table_text,
                                    metadata={
                                        'page': page_num,
                                        'table_index': table_idx,
                                        'is_price_table': is_price_table,
                                        'row_count': len(cleaned) - 1
                                    }
                                ))

            # Compile results
            result['text_elements'] = text_elements

            if all_table_rows and unified_schema:
                result['table_data'] = [{
                    'schema': unified_schema,
                    'rows': all_table_rows
                }]
                result['has_tables'] = True
                result['table_count'] = 1  # Unified into one table
                result['row_count'] = len(all_table_rows)

                logger.info(f"[Feature #216] Hybrid extraction: {len(text_elements)} text elements, "
                           f"{len(all_table_rows)} table rows with schema: {unified_schema[:5]}...")
            else:
                logger.info(f"[Feature #216] Text-only extraction: {len(text_elements)} elements (no valid tables found)")

            return result

        except Exception as e:
            logger.error(f"Error in hybrid PDF extraction: {e}")
            # Fallback to text-only
            result['text_elements'] = self._extract_pdf_structure(file_path)
            return result
